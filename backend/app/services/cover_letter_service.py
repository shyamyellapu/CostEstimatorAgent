"""
Cover Letter Service — orchestrates quotation parsing → AI draft → PDF render pipeline.
REQUIRES quotation upload. Cannot generate without parsed quotation data.
"""
import logging
import json
import io
from typing import Dict, Any, Optional
from pathlib import Path

from docx import Document

from app.config import settings
from app.ai import get_ai_provider
from app.services.document_parser import get_file_text

logger = logging.getLogger(__name__)

from app.services.contract_clauses import CLAUSES, FABRICATION_ONLY_DISCLAIMER

class CoverLetterService:

    def __init__(self):
        self._master_template_text_cache: Optional[str] = None
        self._header_footer_assets_cache: Optional[Dict[str, Any]] = None

    def _master_template_path(self) -> Path:
        project_root = Path(__file__).resolve().parents[3]
        configured = Path(settings.cover_letter_master_template_path)
        if configured.is_absolute():
            return configured
        return project_root / configured

    def _header_footer_docx_path(self) -> Path:
        project_root = Path(__file__).resolve().parents[3]
        configured = Path(settings.cover_letter_header_footer_docx_path)
        if configured.is_absolute():
            return configured
        return project_root / configured

    def _load_header_footer_assets(self) -> Optional[Dict[str, Any]]:
        """Extract header/footer images from the branded DOCX template."""
        if self._header_footer_assets_cache is not None:
            return self._header_footer_assets_cache

        docx_path = self._header_footer_docx_path()
        if not docx_path.exists():
            logger.warning("Header/footer DOCX not found at %s. Using default layout.", docx_path)
            self._header_footer_assets_cache = None
            return None

        try:
            from docx.oxml.ns import qn

            doc = Document(str(docx_path))
            section = doc.sections[0]

            header_blob = None
            footer_blob = None
            header_emu_w = header_emu_h = 0
            footer_emu_w = footer_emu_h = 0

            for rel in section.header.part.rels.values():
                if 'image' in rel.reltype:
                    header_blob = rel.target_part.blob
                    break

            for rel in section.footer.part.rels.values():
                if 'image' in rel.reltype:
                    footer_blob = rel.target_part.blob
                    break

            # Read anchor/inline extent from header XML (EMU units).
            hdr_xml = section.header._element
            for tag in ('wp:anchor', 'wp:inline'):
                for drawing in hdr_xml.iter(qn(tag)):
                    ext = drawing.find(qn('wp:extent'))
                    if ext is not None:
                        header_emu_w = int(ext.get('cx', 0))
                        header_emu_h = int(ext.get('cy', 0))

            ftr_xml = section.footer._element
            for tag in ('wp:anchor', 'wp:inline'):
                for drawing in ftr_xml.iter(qn(tag)):
                    ext = drawing.find(qn('wp:extent'))
                    if ext is not None:
                        footer_emu_w = int(ext.get('cx', 0))
                        footer_emu_h = int(ext.get('cy', 0))

            if not header_blob and not footer_blob:
                logger.warning("No images found in header/footer of %s.", docx_path)
                self._header_footer_assets_cache = None
                return None

            EMU_PER_PT = 12700  # 1 point = 12700 EMU
            assets: Dict[str, Any] = {}
            if header_blob:
                assets["header_png"] = header_blob
                assets["header_w_pt"] = header_emu_w / EMU_PER_PT if header_emu_w else 0
                assets["header_h_pt"] = header_emu_h / EMU_PER_PT if header_emu_h else 0
            if footer_blob:
                assets["footer_png"] = footer_blob
                assets["footer_w_pt"] = footer_emu_w / EMU_PER_PT if footer_emu_w else 0
                assets["footer_h_pt"] = footer_emu_h / EMU_PER_PT if footer_emu_h else 0

            self._header_footer_assets_cache = assets
            return self._header_footer_assets_cache
        except Exception as e:
            logger.warning("Failed to extract header/footer from %s (%s). Using default layout.", docx_path, e)
            self._header_footer_assets_cache = None
            return None

    def load_master_template_text(self) -> str:
        """Load master fabrication template text from DOCX (paragraphs + table cells)."""
        if self._master_template_text_cache:
            return self._master_template_text_cache

        template_path = self._master_template_path()
        if not template_path.exists():
            logger.warning("Master fabrication template not found at %s; falling back to static clauses.", template_path)
            fallback = json.dumps(
                {"clauses": CLAUSES, "disclaimer": FABRICATION_ONLY_DISCLAIMER},
                ensure_ascii=True,
            )
            self._master_template_text_cache = fallback
            return fallback

        doc = Document(str(template_path))
        lines = []

        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if text:
                lines.append(text)

        for table in doc.tables:
            for row in table.rows:
                row_values = []
                for cell in row.cells:
                    cell_text = " ".join(p.text.strip() for p in cell.paragraphs if p.text.strip())
                    if cell_text:
                        row_values.append(cell_text)
                if row_values:
                    lines.append(" | ".join(row_values))

        combined_text = "\n".join(lines)
        payload = json.dumps(
            {
                "master_template_source": str(template_path),
                "disclaimer": FABRICATION_ONLY_DISCLAIMER,
                "master_template_text": combined_text[:30000],
            },
            ensure_ascii=True,
        )
        self._master_template_text_cache = payload
        return payload

    async def parse_quotation(self, file_bytes: bytes, filename: str) -> Dict[str, Any]:
        """Extract structured data from uploaded quotation PDF."""
        text = get_file_text(file_bytes, filename)
        if not text.strip():
            raise ValueError("Could not extract text from quotation file. Ensure it is a text-based PDF.")

        ai = get_ai_provider()
        parsed = await ai.parse_quotation(text)
        logger.info(f"Quotation parsed: client={parsed.get('client')}, ref={parsed.get('reference_number')}")
        return parsed

    async def generate_draft(
        self,
        quotation_data: Dict[str, Any],
        company_info: Optional[Dict[str, str]] = None,
        template_clauses: Optional[str] = None,
    ):
        """Generate cover letter draft from quotation data using AI."""
        if not company_info:
            company_info = {
                "name": settings.company_name,
                "address": settings.company_address,
                "phone": settings.company_phone,
                "email": settings.company_email,
                "website": settings.company_website,
                "signatory_name": settings.signatory_name,
                "signatory_title": settings.signatory_title,
            }

        template_context = template_clauses or self.load_master_template_text()

        ai = get_ai_provider()
        draft = await ai.draft_cover_letter(
            quotation_data=quotation_data,
            template_clauses=template_context,
            company_info=company_info
        )
        return draft

    def render_pdf(self, draft_data: Dict[str, Any], company_info: Dict[str, str]) -> bytes:
        """Render cover letter draft to branded PDF using WeasyPrint."""
        from jinja2 import Environment, FileSystemLoader
        from pathlib import Path

        template_dir = Path(__file__).parent.parent / "templates" / "cover_letter"
        env = Environment(loader=FileSystemLoader(str(template_dir)))

        try:
            template = env.get_template("template.html")
        except Exception:
            # Fallback inline template
            template_str = self._get_inline_template()
            from jinja2 import Template
            template = Template(template_str)

        html = template.render(draft=draft_data, company=company_info)

        try:
            from weasyprint import HTML
            pdf_bytes = HTML(string=html).write_pdf()
            return pdf_bytes
        except Exception as e:
            # WeasyPrint may import successfully but fail at runtime if GTK/Pango/Cairo
            # native libraries are missing on Windows. Fallback to reportlab in all cases.
            logger.warning("WeasyPrint unavailable at runtime (%s). Falling back to reportlab.", e)
            return self._render_with_reportlab(draft_data, company_info)

    def _render_with_reportlab(self, draft: Dict, company: Dict) -> bytes:
        """PDF renderer: C&J Gulf branded layout with letterhead/footer from sample PDF."""
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import cm
        from reportlab.lib import colors
        from reportlab.lib.utils import ImageReader
        from reportlab.platypus import (
            BaseDocTemplate, Frame, PageTemplate, NextPageTemplate,
            Paragraph, Spacer, HRFlowable, Table, TableStyle, Image as RLImage,
        )

        assets = self._load_header_footer_assets()
        page_width, page_height = A4
        left_margin = 2 * cm
        right_margin = 2 * cm

        header_reader = None
        footer_reader = None
        header_height_pt = 0.0
        footer_height_pt = 0.0
        header_draw_w = page_width
        footer_draw_w = page_width
        header_x = 0.0
        footer_x = 0.0

        if assets:
            if "header_png" in assets:
                header_reader = ImageReader(io.BytesIO(assets["header_png"]))
                if assets.get("header_w_pt"):
                    header_draw_w = assets["header_w_pt"]
                    header_height_pt = assets["header_h_pt"]
                else:
                    hw, hh = header_reader.getSize()
                    header_draw_w = page_width
                    header_height_pt = page_width * (hh / hw)
                header_x = (page_width - header_draw_w) / 2

            if "footer_png" in assets:
                footer_reader = ImageReader(io.BytesIO(assets["footer_png"]))
                if assets.get("footer_w_pt"):
                    footer_draw_w = assets["footer_w_pt"]
                    footer_height_pt = assets["footer_h_pt"]
                else:
                    fw, fh = footer_reader.getSize()
                    footer_draw_w = page_width
                    footer_height_pt = page_width * (fh / fw)
                footer_x = (page_width - footer_draw_w) / 2

        # First-page top margin accounts for header image + "To" block + date line.
        to_block_height = 2.6 * cm
        first_top  = header_height_pt + to_block_height + 0.8 * cm
        later_top  = header_height_pt + 0.6 * cm
        bottom_margin = max(2 * cm, footer_height_pt + 0.5 * cm)
        frame_width = page_width - left_margin - right_margin

        first_frame = Frame(
            left_margin, bottom_margin, frame_width,
            page_height - first_top - bottom_margin,
            id='first', leftPadding=0, rightPadding=0, topPadding=0, bottomPadding=0,
        )
        later_frame = Frame(
            left_margin, bottom_margin, frame_width,
            page_height - later_top - bottom_margin,
            id='later', leftPadding=0, rightPadding=0, topPadding=0, bottomPadding=0,
        )

        to_name     = draft.get("to_name", "")
        to_company  = draft.get("to_company", "")
        # Avoid double "M/s." prefix if AI already included it
        display_co = to_company or to_name
        if display_co and not display_co.lower().startswith("m/s"):
            display_co = f"M/s. {display_co}"
        recipient   = display_co
        date_str    = draft.get("date", "")

        def _draw_images(canvas):
            if header_reader:
                canvas.drawImage(header_reader, header_x, page_height - header_height_pt,
                                 width=header_draw_w, height=header_height_pt, mask='auto')
            if footer_reader:
                canvas.drawImage(footer_reader, footer_x, 0,
                                 width=footer_draw_w, height=footer_height_pt, mask='auto')

        def draw_first_page(canvas, _doc):
            _draw_images(canvas)
            canvas.saveState()
            canvas.setFont("Helvetica", 10)
            # "To" block — left aligned, just below header
            y = page_height - header_height_pt - 0.9 * cm
            canvas.drawString(left_margin, y, "To")
            canvas.drawString(left_margin, y - 14, recipient)
            # Date — right aligned at same level
            canvas.drawRightString(page_width - right_margin, y, date_str)
            canvas.restoreState()

        def draw_later_pages(canvas, _doc):
            _draw_images(canvas)

        first_tpl = PageTemplate(id='first', frames=[first_frame], onPage=draw_first_page)
        later_tpl = PageTemplate(id='later', frames=[later_frame], onPage=draw_later_pages)

        buf = io.BytesIO()
        doc = BaseDocTemplate(
            buf, pagesize=A4,
            pageTemplates=[first_tpl, later_tpl],
            leftMargin=left_margin, rightMargin=right_margin,
            topMargin=first_top, bottomMargin=bottom_margin,
        )

        styles = getSampleStyleSheet()
        title_style = ParagraphStyle("CL_Title",
            parent=styles["Heading1"],
            fontSize=16, textColor=colors.HexColor("#1F3864"), spaceAfter=6)
        heading_style = ParagraphStyle("CL_SectionHeading",
            parent=styles["Heading2"],
            fontSize=11, textColor=colors.HexColor("#1F3864"),
            spaceAfter=4, spaceBefore=10)
        body_style = ParagraphStyle("CL_Body",
            parent=styles["Normal"],
            fontSize=10, leading=14, spaceAfter=4)
        bold_style = ParagraphStyle("CL_Bold",
            parent=styles["Normal"],
            fontSize=10, leading=14, fontName="Helvetica-Bold")

        # Switch to later-page template after the first page
        story = [NextPageTemplate('later')]

        # If no image assets, render a text-based company header inline
        if not assets:
            story += [
                Paragraph(company.get("name", ""), title_style),
                Paragraph(company.get("address", ""), body_style),
                Paragraph(f"Tel: {company.get('phone', '')} | Email: {company.get('email', '')}", body_style),
                HRFlowable(width="100%", thickness=2, color=colors.HexColor("#1F3864")),
                Spacer(1, 0.3 * cm),
                # Date (right) + To block (left) as a table
                Table(
                    [[Paragraph("To", body_style),
                      Paragraph(date_str, body_style)],
                     [Paragraph(f"<b>{recipient}</b>", bold_style),
                      Paragraph("", body_style)]],
                    colWidths=[frame_width * 0.6, frame_width * 0.4],
                    style=TableStyle([
                        ('ALIGN', (1, 0), (1, -1), 'RIGHT'),
                        ('VALIGN', (0, 0), (-1, -1), 'TOP'),
                        ('LEFTPADDING', (0, 0), (-1, -1), 0),
                        ('RIGHTPADDING', (0, 0), (-1, -1), 0),
                    ]),
                ),
                Spacer(1, 0.4 * cm),
            ]

        # Reference line (no Subject per requirement #5)
        story.append(Paragraph(f"<b>Our Ref:</b> {draft.get('reference', '')}", body_style))
        story.append(Spacer(1, 0.4 * cm))

        # Gender-aware salutation (#1)
        salutation = draft.get("salutation") or "Dear Sir/Madam,"
        story.append(Paragraph(salutation, body_style))
        story.append(Spacer(1, 0.3 * cm))

        # Phrases to strip from AI-generated scope content (#2)
        _STRIP = [
            "engineering and design responsibility, preparation of fabrication drawings, site installation, commissioning,",
            ", preparation of fabrication drawings, site installation, commissioning,",
            "and design responsibility, preparation of fabrication drawings, site installation, commissioning,",
        ]

        # Fixed Validity + closing text (#8)
        _VALIDITY_TEXT = (
            "This offer shall remain valid for 30 days from the date of submission. Upon award, the Purchase "
            "Order, approved scope, agreed drawings, and this techno-commercial proposal shall together "
            "form the basis of the contract.\n\n"
            "We trust the above clarifies our scope, responsibilities, and commercial basis. Should you "
            "require any clarification or additional information, please feel free to contact us. We look "
            "forward to your valued order and to executing this project to your satisfaction.\n\n"
            "Should you require any clarification or alignment discussion on the above, we remain available "
            "for coordination.\n\n"
            f"We appreciate the opportunity to engage with {to_company or to_name} and look forward to your "
            "confirmation and continued professional association."
        )

        for section in draft.get("sections", []):
            story.append(Paragraph(section.get("title", ""), heading_style))
            content = section.get("content", "")
            # Strip disallowed exclusion phrases (#2)
            for phrase in _STRIP:
                content = content.replace(phrase, "")
            # Replace validity section with prescribed text (#8)
            if section.get("section_id") == "validity":
                content = _VALIDITY_TEXT
            for para in content.split("\n"):
                if para.strip():
                    story.append(Paragraph(para.strip(), body_style))

        # Closing lines
        story.append(Spacer(1, 0.8 * cm))
        story.append(Paragraph("Yours faithfully,", body_style))
        story.append(Paragraph("For C&amp;J Gulf Equipment Manufacturing LLC.", body_style))

        # --- Signature images + names ---
        project_root = Path(__file__).resolve().parents[3]
        sig_dir = project_root / "ReferenceFiles"
        sig_h = 38  # pt — tuned via signature_editor

        def _sig_img(filename):
            p = sig_dir / filename
            if not p.exists():
                return Spacer(1, sig_h)
            reader = ImageReader(str(p))
            iw, ih = reader.getSize()
            img = RLImage(str(p), width=sig_h * (iw / ih), height=sig_h)
            img.hAlign = 'LEFT'
            return img

        bilal_img  = _sig_img("BilalAhmed Signature.jpg")
        datta_img  = _sig_img("Datta C.Sawant Signature.jpg")
        subash_img = _sig_img("subash Valrani Signature.jpg")

        # Bilal: left-aligned signature, name/title
        story.append(Spacer(1, 10))  # BILAL_GAP
        story.append(bilal_img)
        story.append(Paragraph("<b>Bilal Ahmed</b>", bold_style))
        story.append(Paragraph("Cost &amp; Estimation Engineer.", body_style))
        story.append(Spacer(1, 14))  # DATTA_GAP

        # Datta + Subash side-by-side
        half = frame_width / 2
        sig_table = Table(
            [
                [datta_img, subash_img],
                [Paragraph("<b>Datta C. Sawant</b>", bold_style),
                 Paragraph("<b>Subash Valrani</b>", bold_style)],
                [Paragraph("Sr. Mechanical Engineer", body_style),
                 Paragraph("Business Unit Head", body_style)],
            ],
            colWidths=[half, half],
        )
        sig_table.setStyle(TableStyle([
            ('VALIGN',        (0, 0), (-1, -1), 'TOP'),
            ('LEFTPADDING',   (0, 0), (-1, -1), 0),
            ('RIGHTPADDING',  (0, 0), (-1, -1), 0),
            ('TOPPADDING',    (0, 0), (-1, -1), 1),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 1),
        ]))
        story.append(sig_table)

        doc.build(story)
        buf.seek(0)
        return buf.read()

    def _get_inline_template(self) -> str:
        return """<!DOCTYPE html>
<html>
<head>
<meta charset="utf-8">
<style>
  body { font-family: 'Arial', sans-serif; font-size: 10pt; color: #333; margin: 0; padding: 0; }
  .header { background: #1F3864; color: white; padding: 20px 30px; }
  .header h1 { margin: 0; font-size: 18pt; }
  .header p { margin: 4px 0; font-size: 9pt; }
  .divider { border-top: 3px solid #1F3864; margin: 0; }
  .content { padding: 25px 35px; }
  .meta { margin: 15px 0; }
  .section-heading { color: #1F3864; font-size: 11pt; font-weight: bold;
                     border-bottom: 1px solid #D9E1F2; padding-bottom: 3px; margin-top: 15px; }
  .section-content { margin: 8px 0 0 0; line-height: 1.5; font-size: 10pt; }
  .closing { margin-top: 25px; }
  .signatory { margin-top: 40px; }
  .sig-line { border-top: 1px solid #333; width: 200px; margin: 5px 0; }
</style>
</head>
<body>
<div class="header">
  <h1>{{ company.name }}</h1>
  <p>{{ company.address }}</p>
  <p>Tel: {{ company.phone }} | Email: {{ company.email }}</p>
</div>
<div class="divider"></div>
<div class="content">
  <div class="meta">
    <p><strong>Date:</strong> {{ draft.date }}</p>
    <p><strong>To:</strong> {{ draft.to_name }}, {{ draft.to_company }}</p>
    <p><strong>Ref:</strong> {{ draft.reference }}</p>
    <p><strong>Subject:</strong> {{ draft.subject }}</p>
  </div>
  <p>Dear Sir/Madam,</p>
  {% for section in draft.sections %}
  <div class="section-heading">{{ section.title }}</div>
  <div class="section-content">{{ section.content | replace('\n', '<br>') }}</div>
  {% endfor %}
  <div class="closing">
    <p>{{ draft.closing }}</p>
    <p>Yours faithfully,</p>
  </div>
  <div class="signatory">
    <div class="sig-line"></div>
    <p><strong>{{ draft.signatory_name }}</strong></p>
    <p>{{ draft.signatory_title }}</p>
    <p>{{ company.name }}</p>
  </div>
</div>
</body>
</html>"""


cover_letter_service = CoverLetterService()
