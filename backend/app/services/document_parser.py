"""
Document Parser — extracts text from PDF, DOCX, Excel, and image files.
Text extraction is pre-processing before AI calls.
"""
import io
import logging
from pathlib import Path
from typing import Optional

logger = logging.getLogger(__name__)


def extract_text_from_pdf(file_bytes: bytes) -> str:
    """Extract all text from a PDF file using PyMuPDF with a pypdf fallback."""
    try:
        import fitz
        doc = fitz.open(stream=file_bytes, filetype="pdf")
        texts = []
        for page in doc:
            texts.append(page.get_text())
        doc.close()
        return "\n".join(texts)
    except Exception as e:
        logger.warning(f"PyMuPDF PDF text extraction failed: {e}")

    try:
        from pypdf import PdfReader
        reader = PdfReader(io.BytesIO(file_bytes))
        texts = []
        for page in reader.pages:
            page_text = page.extract_text() or ""
            if page_text.strip():
                texts.append(page_text)
        return "\n".join(texts)
    except Exception as fallback_error:
        logger.error(f"PDF text extraction error: {fallback_error}")
        return ""


def pdf_page_to_image(file_bytes: bytes, page_num: int = 0) -> bytes:
    """Convert a specific PDF page to PNG bytes for Vision processing."""
    try:
        import fitz
        doc = fitz.open(stream=file_bytes, filetype="pdf")
        if page_num >= len(doc):
            return b""
        page = doc[page_num]
        pix = page.get_pixmap(matrix=fitz.Matrix(1.5, 1.5))  # 1.5x scale balances quality and size
        img_bytes = pix.tobytes("png")
        doc.close()
        return img_bytes
    except Exception as e:
        logger.error(f"PDF to Image conversion error: {e}")
        return b""


def pdf_to_images(file_bytes: bytes, max_pages: int = 3, scale: float = 1.5) -> list[bytes]:
    """Convert up to max_pages of a PDF into PNG bytes for multi-page vision extraction."""
    try:
        import fitz
        doc = fitz.open(stream=file_bytes, filetype="pdf")
        page_count = min(len(doc), max_pages)
        images = []
        for page_num in range(page_count):
            page = doc[page_num]
            pix = page.get_pixmap(matrix=fitz.Matrix(scale, scale))
            images.append(pix.tobytes("png"))
        doc.close()
        return images
    except Exception as e:
        logger.error(f"PDF to Images conversion error: {e}")
        return []


def extract_text_from_docx(file_bytes: bytes) -> str:
    """Extract text from DOCX."""
    try:
        from docx import Document
        doc = Document(io.BytesIO(file_bytes))
        paras = [p.text for p in doc.paragraphs if p.text.strip()]
        # Also extract tables
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells)
                if row_text.strip():
                    paras.append(row_text)
        return "\n".join(paras)
    except Exception as e:
        logger.error(f"DOCX text extraction error: {e}")
        return ""


def extract_text_from_excel(file_bytes: bytes) -> str:
    """Extract text/data from Excel as tabular text."""
    try:
        import openpyxl
        wb = openpyxl.load_workbook(io.BytesIO(file_bytes), data_only=True)
        rows = []
        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            rows.append(f"=== Sheet: {sheet_name} ===")
            for row in ws.iter_rows(values_only=True):
                row_str = " | ".join(str(c) if c is not None else "" for c in row)
                if row_str.strip(" |"):
                    rows.append(row_str)
        return "\n".join(rows)
    except Exception as e:
        logger.error(f"Excel text extraction error: {e}")
        return ""


def get_file_text(file_bytes: bytes, filename: str, mime_type: Optional[str] = None) -> str:
    """Dispatcher — extract text based on file type."""
    ext = Path(filename).suffix.lower()

    if ext == ".pdf" or (mime_type and "pdf" in mime_type):
        return extract_text_from_pdf(file_bytes)
    elif ext in (".docx", ".doc") or (mime_type and "word" in mime_type):
        return extract_text_from_docx(file_bytes)
    elif ext in (".xlsx", ".xls") or (mime_type and "spreadsheet" in mime_type):
        return extract_text_from_excel(file_bytes)
    elif ext in (".txt", ".csv"):
        return file_bytes.decode("utf-8", errors="replace")
    else:
        # Try plain text
        try:
            return file_bytes.decode("utf-8", errors="replace")
        except Exception:
            return ""


def is_image_file(filename: str) -> bool:
    ext = Path(filename).suffix.lower()
    return ext in (".png", ".jpg", ".jpeg", ".gif", ".bmp", ".tiff", ".webp")
